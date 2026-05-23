#!/opt/homebrew/bin/python3.13
"""Generate the 4 per-manager discipline-feedback Word documents for v24.3.

One docx per discipline manager. Each pre-populated with Cowork's findings
(verdict, evidence_visual, evidence_pages, compliance_note), with hand-fillable
checkboxes for the manager to confirm/override during the meeting.

Output: templates/v24_3_discipline_feedback_{slug}.docx
Run: /opt/homebrew/bin/python3.13 scripts/gen_discipline_feedback_templates.py
"""
from __future__ import annotations

import json
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Mm, RGBColor

ROOT = Path(__file__).resolve().parent.parent
AUDIT_RESULTS = ROOT / "audit_outputs" / "407-1048248" / "v24.3" / "audit_results.json"
SUBMISSION_META = ROOT / "projects" / "407-1048248" / "submissions" / "v24.3" / "metadata.json"
LOGO = ROOT / "assets" / "nessziona_logo.png"
OUT_DIR = ROOT / "templates"

PROJECT_NAME_HE = "מתחם הטייסים-ההסתדרות"
TAVA = "407-1048248"
SUBMISSION_VERSION = "24.3"

# Discipline → Hebrew name (matches the engine's renderer)
DISCIPLINE_NAME_HE = {
    "shafa":    'שפ"ע — אשפה ופינוי פסולת',
    "gardens":  "גנים ונוף",
    "infra":    "תשתיות",
    "fire":     "רחבות כיבוי",
    "drainage": "ניקוז וחלחול",
    "roofs":    "גגות וחזית חמישית",
    "arch":     "אדריכלות וחזיתות",
    "balcony":  "מרפסות",
    "laundry":  "מסתורי כביסה",
    "env":      "הנחיות סביבתיות",
}

# Verdict → Hebrew display label (engine taxonomy → manager-friendly)
VERDICT_HE = {
    "pass":              "תקין",
    "fail":              "נדרש תיקון",
    "requires_review":   "דורש בירור",
    "not_submitted":     "לא הוגש",
    "unevaluable":       "לא ניתן לבדיקה",
}

# Engine taxonomy that managers can choose between (3 checkboxes per rule)
MANAGER_VERDICT_OPTIONS = [
    ("pass",            "תקין"),
    ("fail",            "נדרש תיקון"),
    ("requires_review", "דורש בירור"),
]

# 4 managers → list of disciplines they cover
MANAGER_MAPPING = [
    {
        "slug": "shafa_fire_env",
        "title_he": 'שפ"ע',
        "disciplines": ["shafa", "fire", "env"],
        "subtitle_he": 'אשפה ופינוי פסולת · רחבות כיבוי · הנחיות סביבתיות',
    },
    {
        "slug": "gardens_roofs",
        "title_he": "גינון",
        "disciplines": ["gardens", "roofs"],
        "subtitle_he": "גנים ונוף · גגות וחזית חמישית",
    },
    {
        "slug": "infra_drainage",
        "title_he": "תשתיות",
        "disciplines": ["infra", "drainage"],
        "subtitle_he": "תשתיות · ניקוז וחלחול",
    },
    {
        "slug": "arch_balcony_laundry",
        "title_he": "אדריכלית העיר",
        "disciplines": ["arch", "balcony", "laundry"],
        "subtitle_he": "אדריכלות וחזיתות · מרפסות · מסתורי כביסה",
    },
]


# ─────────────────────────────────────────────────────────────────────────────
# RTL / styling helpers (python-docx high-level API has gaps for RTL)
# ─────────────────────────────────────────────────────────────────────────────

def _set_rtl(paragraph) -> None:
    """Mark paragraph as RTL — sets <w:bidi/> in pPr."""
    pPr = paragraph._p.get_or_add_pPr()
    bidi = OxmlElement("w:bidi")
    pPr.append(bidi)


def _set_run_rtl(run) -> None:
    """Mark run as containing RTL text + apply Hebrew font to all script slots."""
    rPr = run._element.get_or_add_rPr()
    rtl = OxmlElement("w:rtl")
    rPr.append(rtl)
    # Apply Heebo to all 4 font script slots so Word doesn't fall back to a
    # different font for Hebrew characters.
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    for slot in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(slot), "Heebo")


def _add_para(doc_or_cell, text: str, *,
              bold: bool = False, italic: bool = False, size_pt: float = 11,
              color: tuple[int, int, int] | None = None,
              align=WD_ALIGN_PARAGRAPH.RIGHT,
              space_before_pt: float = 0, space_after_pt: float = 0):
    p = doc_or_cell.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before_pt)
    p.paragraph_format.space_after = Pt(space_after_pt)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.3
    _set_rtl(p)
    if text:
        run = p.add_run(text)
        run.bold = bold
        run.italic = italic
        run.font.size = Pt(size_pt)
        if color is not None:
            run.font.color.rgb = RGBColor(*color)
        _set_run_rtl(run)
    return p


def _add_heading(doc, text: str, level: int):
    sizes = {1: 22, 2: 16, 3: 13}
    colors = {1: (0, 80, 48), 2: (0, 80, 48), 3: (60, 60, 60)}
    space = {1: (10, 6), 2: (14, 4), 3: (10, 3)}
    sb, sa = space[level]
    return _add_para(doc, text, bold=True, size_pt=sizes[level], color=colors[level],
                     space_before_pt=sb, space_after_pt=sa)


def _add_rule(doc):
    """Horizontal divider."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "C0C0C0")
    pbdr.append(bottom)
    pPr.append(pbdr)
    p.paragraph_format.space_after = Pt(4)


def _add_checkbox_line(doc, options: list[tuple[str, str]]):
    """Add a single paragraph with ☐ option-label  ☐ option-label  ☐ option-label."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    _set_rtl(p)
    for i, (_, label) in enumerate(options):
        sep = "        " if i > 0 else ""
        run = p.add_run(f"{sep}☐  {label}")
        run.font.size = Pt(12)
        _set_run_rtl(run)
    return p


def _add_text_field_line(doc, label: str, height_pts: float = 36):
    """Add a labeled empty 'fill-in' area: label on top, blank underlined-bottom area."""
    _add_para(doc, label, bold=True, size_pt=10.5,
              color=(110, 110, 110), space_before_pt=4, space_after_pt=2)
    # bordered empty paragraph for handwriting
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    pPr = p._p.get_or_add_pPr()
    pbdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "B0B0B0")
    pbdr.append(bottom)
    pPr.append(pbdr)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    # Stack 2 lines worth of vertical space for handwriting
    for _ in range(2):
        p2 = doc.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_rtl(p2)
        pPr2 = p2._p.get_or_add_pPr()
        pbdr2 = OxmlElement("w:pBdr")
        bottom2 = OxmlElement("w:bottom")
        bottom2.set(qn("w:val"), "single")
        bottom2.set(qn("w:sz"), "6")
        bottom2.set(qn("w:space"), "1")
        bottom2.set(qn("w:color"), "B0B0B0")
        pbdr2.append(bottom2)
        pPr2.append(pbdr2)
        p2.paragraph_format.space_before = Pt(8)
        p2.paragraph_format.space_after = Pt(0)


def _add_page_break(doc):
    p = doc.add_paragraph()
    _set_rtl(p)
    p.add_run().add_break(WD_BREAK.PAGE)


def _set_default_font(doc):
    """Set Heebo as the document's default font (so any unstyled text uses it)."""
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Heebo"
    normal.font.size = Pt(11)
    # Set East-Asian + complex-script slots so Word doesn't substitute fonts
    rpr = normal.element.get_or_add_rPr()
    rFonts = rpr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rpr.append(rFonts)
    for slot in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rFonts.set(qn(slot), "Heebo")
    # Section page margins
    section = doc.sections[0]
    section.top_margin = Mm(20)
    section.bottom_margin = Mm(20)
    section.left_margin = Mm(22)
    section.right_margin = Mm(22)
    # RTL section
    sectPr = section._sectPr
    bidi = OxmlElement("w:bidi")
    sectPr.append(bidi)


# ─────────────────────────────────────────────────────────────────────────────
# Page builders
# ─────────────────────────────────────────────────────────────────────────────

def _build_cover(doc, manager: dict):
    # Logo
    p_logo = doc.add_paragraph()
    p_logo.alignment = WD_ALIGN_PARAGRAPH.LEFT  # logo on the LEFT of an RTL page = visual right
    p_logo.add_run().add_picture(str(LOGO), height=Mm(28))
    p_logo.paragraph_format.space_after = Pt(2)

    _add_para(doc, "מינהלת ההתחדשות העירונית · עיריית נס ציונה",
              size_pt=11, color=(110, 110, 110), space_before_pt=0, space_after_pt=2)
    _add_para(doc, "סקירת תכנית עיצוב — משוב דיסציפלינה",
              bold=True, size_pt=14, color=(0, 80, 48),
              space_before_pt=0, space_after_pt=16)

    _add_para(doc, f"דיסציפלינה: {manager['title_he']}",
              bold=True, size_pt=26, color=(0, 80, 48),
              space_before_pt=0, space_after_pt=4)
    _add_para(doc, manager["subtitle_he"],
              size_pt=13, color=(70, 70, 70), space_after_pt=24)

    _add_rule(doc)

    # Reference info
    _add_para(doc, "פרטי ההגשה", bold=True, size_pt=12, color=(0, 80, 48),
              space_before_pt=10, space_after_pt=4)
    _add_para(doc, f'פרויקט:  {PROJECT_NAME_HE}', size_pt=11, space_after_pt=2)
    _add_para(doc, f'תכנית סטטוטורית:  תב"ע {TAVA}', size_pt=11, space_after_pt=2)
    _add_para(doc, f"גרסת תכנית עיצוב:  {SUBMISSION_VERSION}", size_pt=11, space_after_pt=2)
    _add_para(doc, "תאריך הפגישה:  __________________________", size_pt=11, space_after_pt=2)
    _add_para(doc, "מנהל הדיסציפלינה:  __________________________", size_pt=11, space_after_pt=12)

    _add_para(doc, 'הערה: תכנית העיצוב המלאה (PDF גרסה 24.3) צורפה למסמך זה לעיון.',
              italic=True, size_pt=10, color=(110, 110, 110), space_after_pt=12)

    _add_rule(doc)

    # Instructions
    _add_para(doc, "כיצד למלא את המסמך", bold=True, size_pt=12, color=(0, 80, 48),
              space_before_pt=10, space_after_pt=4)
    instructions = [
        f'כל סעיף בדיסציפלינה {manager["title_he"]} מציג את הממצא הנוכחי של המנוע האוטומטי '
        '(וריפיקציה ויזואלית של Cowork) — שם הכלל, הקביעה הנוכחית, התיאור הויזואלי, ההפניות לעמודים בהגשה, וההערה האוטומטית.',
        'לכל סעיף יש שלוש תיבות לסימון: ☐ תקין · ☐ נדרש תיקון · ☐ דורש בירור. '
        'יש לסמן את הקביעה לאחר דיון. אם מסכימים עם הקביעה הנוכחית — סמן/י אותה שוב לאישור.',
        'בשדה "הערה / נימוק" יש להוסיף הסבר קצר: לפי איזו הוראה (תב"ע / חוברת הנחיות / תקן לאומי) ניתנה הקביעה, או מהי הדרישה התכנונית מהאדריכל.',
        'בסוף המסמך יש סעיף "ממצאים נוספים" — לתיעוד בעיות שלא נכללו בבדיקה האוטומטית.',
        'חוות הדעת הסופית של המינהלת תשלב את המשוב שלך עם ממצאי המנוע, ותועבר לאדריכל.',
    ]
    for i, instr in enumerate(instructions, 1):
        _add_para(doc, f"{i}.  {instr}", size_pt=11, space_after_pt=4)

    _add_para(doc, "לשאלות:  מהנדס/ת המינהלת · מזכירות המינהלת",
              size_pt=10, italic=True, color=(110, 110, 110),
              space_before_pt=10, space_after_pt=2)

    _add_page_break(doc)


def _build_rule_block(doc, rule: dict):
    """One block per rule: title, current findings, manager input."""
    title = rule.get("rule_name_he", rule.get("rule_code", ""))
    verdict = rule.get("verdict", "")
    verdict_he = VERDICT_HE.get(verdict, "—")
    evidence_pages = rule.get("evidence_pages") or rule.get("evidence", {}).get("evidence_pages") or []
    evidence_visual = (rule.get("evidence_visual") or rule.get("evidence", {}).get("evidence_visual") or "").strip()
    compliance_note = (rule.get("compliance_note") or rule.get("evidence", {}).get("compliance_note") or "").strip()

    # rule title
    _add_para(doc, title, bold=True, size_pt=13, color=(0, 80, 48),
              space_before_pt=14, space_after_pt=2)

    # current verdict line + page refs
    pages_str = f"  (עמ' {', '.join(str(p) for p in evidence_pages)})" if evidence_pages else ""
    _add_para(doc, f"קביעה נוכחית של המנוע:  {verdict_he}{pages_str}",
              size_pt=11, color=(70, 70, 70), space_after_pt=4)

    # evidence visual
    if evidence_visual:
        _add_para(doc, "תיאור ויזואלי מההגשה:", bold=True, size_pt=10.5,
                  color=(110, 110, 110), space_before_pt=2, space_after_pt=1)
        _add_para(doc, evidence_visual, italic=True, size_pt=10.5,
                  color=(50, 50, 50), space_after_pt=4)

    # compliance note
    if compliance_note:
        _add_para(doc, "הערת המנוע:", bold=True, size_pt=10.5,
                  color=(110, 110, 110), space_before_pt=2, space_after_pt=1)
        _add_para(doc, compliance_note, size_pt=11, color=(50, 50, 50), space_after_pt=4)

    # manager input: 3 checkboxes
    _add_para(doc, "קביעת המנהל/ת:", bold=True, size_pt=11,
              color=(0, 80, 48), space_before_pt=6, space_after_pt=2)
    _add_checkbox_line(doc, MANAGER_VERDICT_OPTIONS)

    # free-text reason
    _add_text_field_line(doc, "הערה / נימוק:")

    _add_rule(doc)


def _build_additional_findings(doc):
    _add_page_break(doc)
    _add_heading(doc, "ממצאים נוספים שלא בעצימת הבדיקה האוטומטית", level=2)
    _add_para(doc, "לתיעוד בעיות תכנוניות שעלו בדיון ולא נכללו בכללים האוטומטיים. "
                   "כל ממצא חדש מקבל כותרת קצרה, קביעה, ונימוק.",
              italic=True, size_pt=10.5, color=(110, 110, 110), space_after_pt=10)

    for i in range(1, 6):
        _add_para(doc, f"ממצא #{i}", bold=True, size_pt=12, color=(0, 80, 48),
                  space_before_pt=10, space_after_pt=2)
        _add_text_field_line(doc, "כותרת קצרה:")
        _add_para(doc, "קביעה:", bold=True, size_pt=11,
                  color=(0, 80, 48), space_before_pt=4, space_after_pt=2)
        _add_checkbox_line(doc, MANAGER_VERDICT_OPTIONS)
        _add_text_field_line(doc, "תיאור / נימוק:")
        _add_rule(doc)


def _build_signature_block(doc):
    _add_page_break(doc)
    _add_heading(doc, "אישור וחתימה", level=2)
    _add_para(doc, "אני מאשר/ת שהממצאים והקביעות במסמך זה מבטאים את עמדתי המקצועית "
                   "כמנהל/ת הדיסציפלינה.",
              size_pt=11, space_after_pt=20)
    _add_text_field_line(doc, "שם:")
    _add_text_field_line(doc, "תפקיד:")
    _add_text_field_line(doc, "תאריך:")
    _add_text_field_line(doc, "חתימה:")


# ─────────────────────────────────────────────────────────────────────────────
# Main
# ─────────────────────────────────────────────────────────────────────────────

def main() -> int:
    if not AUDIT_RESULTS.exists():
        print(f"ERROR: audit_results.json not found at {AUDIT_RESULTS}", file=sys.stderr)
        return 1

    audit = json.loads(AUDIT_RESULTS.read_text(encoding="utf-8"))
    discipline_results = audit.get("disciplines", [])

    OUT_DIR.mkdir(parents=True, exist_ok=True)

    for manager in MANAGER_MAPPING:
        doc = Document()
        _set_default_font(doc)
        _build_cover(doc, manager)

        # Group rules by discipline, in the manager's assigned order
        for disc_code in manager["disciplines"]:
            disc_rules = [r for r in discipline_results if r.get("discipline") == disc_code]
            if not disc_rules:
                continue
            disc_name = DISCIPLINE_NAME_HE.get(disc_code, disc_code)
            _add_heading(doc, disc_name, level=1)
            for rule in sorted(disc_rules, key=lambda x: x.get("rule_code", "")):
                _build_rule_block(doc, rule)

        _build_additional_findings(doc)
        _build_signature_block(doc)

        out_path = OUT_DIR / f"v24_3_discipline_feedback_{manager['slug']}.docx"
        doc.save(str(out_path))
        rule_count = sum(1 for r in discipline_results if r.get("discipline") in manager["disciplines"])
        print(f"wrote {out_path.name}  ({rule_count} rules across "
              f"{len(manager['disciplines'])} disciplines)")

    return 0


if __name__ == "__main__":
    raise SystemExit(main())
